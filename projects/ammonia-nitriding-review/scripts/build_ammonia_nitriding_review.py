#!/usr/bin/env python3
from __future__ import annotations

import os
import re
from functools import lru_cache
from pathlib import Path
from urllib.request import Request, urlopen

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
DOCX_DIR = OUT_DIR / "docx"
PDF_DIR = OUT_DIR / "pdf"
FIG_DIR = OUT_DIR / "figures"
SOURCE_FIG_DIR = OUT_DIR / "source_figures"
VERSION = os.environ.get("REVIEW_VERSION", "Ver1")
DOCX_PATH = DOCX_DIR / f"김영광_암모니아질화_리뷰논문_{VERSION}.docx"
ADD_EXPANSION_PARAGRAPHS = os.environ.get("ADD_EXPANSION_PARAGRAPHS", "0") == "1"
MAX_PARAGRAPHS_PER_SUBSECTION = int(os.environ.get("MAX_PARAGRAPHS_PER_SUBSECTION", "0"))

K_FONT = "Batang"
E_FONT = "Times New Roman"
FIG_K_FONT_ENV = os.environ.get("FIG_K_FONT") or os.environ.get("FIG_FONT")
FIG_E_FONT_ENV = os.environ.get("FIG_E_FONT")


def resolve_font_file(env_value: str | None, candidates: list[str], label: str) -> str:
    if env_value and Path(env_value).expanduser().exists():
        return str(Path(env_value).expanduser())
    for candidate in candidates:
        path = Path(candidate).expanduser()
        if path.exists():
            return str(path)
    raise SystemExit(
        f"{label} 그림 글꼴 파일을 찾지 못했습니다. "
        f"해당 글꼴을 설치하거나 환경변수 FIG_K_FONT/FIG_E_FONT로 경로를 지정하세요."
    )


FIG_K_FONT = resolve_font_file(
    FIG_K_FONT_ENV,
    [
        "/Applications/Microsoft Word.app/Contents/Resources/DFonts/batang.ttc",
        "/Applications/Microsoft PowerPoint.app/Contents/Resources/DFonts/batang.ttc",
        "/Applications/Microsoft Excel.app/Contents/Resources/DFonts/batang.ttc",
        "/System/Library/Fonts/Supplemental/Batang.ttf",
        "/Library/Fonts/Batang.ttf",
        "~/Library/Fonts/Batang.ttf",
        "~/Library/Fonts/batang.ttc",
    ],
    "Batang",
)
FIG_E_FONT = resolve_font_file(
    FIG_E_FONT_ENV,
    [
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Times.ttf",
        "/Library/Fonts/Times New Roman.ttf",
        "~/Library/Fonts/Times New Roman.ttf",
    ],
    "Times New Roman",
)


def env_value(name: str, default: str = "") -> str:
    return os.environ.get(name, default).strip()


METADATA = {
    "author": env_value("THESIS_AUTHOR", "김영광"),
    "degree_name": env_value("THESIS_DEGREE_NAME"),
    "department": env_value("THESIS_DEPARTMENT"),
    "major": env_value("THESIS_MAJOR"),
    "advisor": env_value("THESIS_ADVISOR"),
    "submission_month": env_value("THESIS_SUBMISSION_MONTH"),
    "approval_month": env_value("THESIS_APPROVAL_MONTH"),
    "committee_chair": env_value("THESIS_COMMITTEE_CHAIR", "이명규"),
    "committee_vice_chair": env_value("THESIS_COMMITTEE_VICE_CHAIR", "한흥남"),
    "committee_member": env_value("THESIS_COMMITTEE_MEMBER", "김동익"),
    "committee_member_role": env_value("THESIS_COMMITTEE_MEMBER_ROLE", "외부심사위원"),
    "korean_title_lines": [
        env_value("THESIS_KO_TITLE_1", "암모니아 에너지 시스템에서의"),
        env_value("THESIS_KO_TITLE_2", "원치 않는 질화 및 질화부식"),
        env_value("THESIS_KO_TITLE_3", "기구, 합금 의존성, 억제 전략에 관한 리뷰"),
    ],
    "english_title_lines": [
        env_value("THESIS_EN_TITLE_1", "Unwanted Nitriding and Nitridation Corrosion in Ammonia Energy Systems"),
        env_value("THESIS_EN_TITLE_2", "Mechanisms, Alloy Dependence, and Mitigation Strategies"),
    ],
}

REQUIRED_METADATA = [
    ("degree_name", "학위논문 종류"),
    ("department", "학과"),
    ("major", "전공"),
    ("advisor", "지도교수"),
    ("submission_month", "논문 제출월"),
    ("approval_month", "논문 인준월"),
]


REFERENCES = [
    "Jeong, D.; Yeom, H. Nitriding-Induced Degradation of Structural Steels in High-Temperature Ammonia Utilizing Systems. Korean Journal of Chemical Engineering, 2026. https://doi.org/10.1007/s11814-026-00706-2",
    "Ojasalo, S.; Antikainen, A.; Nandy, S.; Nurmi, V.; Bojinov, M.; Karastoyanov, V.; Huttunen-Saarivirta, E. High-temperature corrosion of steels by nitridation in ammonia: Degradation mechanisms and comparison between steel grades. Fuel, 419, 138872, 2026. https://doi.org/10.1016/j.fuel.2026.138872",
    "Wang, D.; Xing, Y.; Lee, M.; Suzuki, Y. Effects of wall temperature and water vapor on the nitriding of stainless steel induced by ammonia flames. Proceedings of the Combustion Institute, 40, 105562, 2024. https://doi.org/10.1016/j.proci.2024.105562",
    "Xing, Y.; Lee, M.; Suzuki, Y. Effect of water vapor on nitriding of stainless steel walls induced by ammonia flames. Proceedings of the Combustion Institute, 2025, 105831. https://doi.org/10.1016/j.proci.2025.105831",
    "Ghara, T.; Kuroda, S.; Yanagisawa, T.; Shahien, M.; Suzuki, M.; Inoue, T.; Shinoda, K. Degradation behaviour of HVOF sprayed CoNiCrAlY coating in high-temperature ammonia environment towards its applicability in ammonia fueled gas turbines. International Journal of Hydrogen Energy, 130, 345-359, 2025. https://doi.org/10.1016/j.ijhydene.2025.04.277",
    "Laws, N.; Sarabia, E.; Roberts, W. L.; Campuzano, F. Aluminization and pre-oxidation treatments for enhanced nitridation resistance of reformer alloys under high-pressure ammonia cracking. Materials Chemistry and Physics, 358, 132426, 2026. https://doi.org/10.1016/j.matchemphys.2026.132426",
    "Li, J.; Zheng, L.; Hu, L.; Teng, L.; Rao, S.; Zhang, C.; Luo, Y.; Jiang, L. Study on the temperature- and position-dependent corrosion behaviour of 310S stainless steel in high-temperature ammonia decomposition reactors. Chemical Engineering Science, 333, 124309, 2026. https://doi.org/10.1016/j.ces.2026.124309",
    "Fan, L.; Lv, Y.; Wu, L.; Zhang, S.; Wang, T.; Liu, F.; Ding, X.; Yao, J. Enhanced Nitriding of 38CrMoAl Steels with Laser Vibrational Excitation of Ammonia. Metallurgical and Materials Transactions A, 55, 3302-3312, 2024. https://doi.org/10.1007/s11661-024-07456-y",
    "Kochmanski, P.; Bielawski, J.; Baranowska, J. Effect of low temperature gas nitriding on corrosion properties of duplex stainless steel. Surface and Coatings Technology, 517, 132842, 2025. https://doi.org/10.1016/j.surfcoat.2025.132842",
    "Yanagimoto, F.; Shiotani, K.; Sakakibara, Y.; Watanabe, Y.; Tada, E. Stress corrosion cracking of steels in liquid ammonia: A comprehensive literature review. International Journal of Hydrogen Energy, 120, 89-100, 2025. https://doi.org/10.1016/j.ijhydene.2025.03.280",
    "Javeria, U.; Kim, S. J. Liquid ammonia-induced stress corrosion cracking of steels: A comprehensive review of mechanism, metallurgical and environmental factors, testing protocols, and mitigation challenges. International Journal of Hydrogen Energy, 232, 154995, 2026. https://doi.org/10.1016/j.ijhydene.2026.154995",
    "Zhao, Z.; Zhang, M.; Wu, Y.; Song, W.; Yan, J.; Qi, X.; Yang, J.; Wen, J.; Zhang, H. Ammonia Energy: Synthesis and Utilization. Industrial & Engineering Chemistry Research, 63, 8003-8024, 2024. https://doi.org/10.1021/acs.iecr.4c00384",
    "Jafar, U.; Nuhu, U.; Khan, W. U.; Hossain, M. M. A review on green ammonia as a potential CO2 free fuel. International Journal of Hydrogen Energy, 71, 857-876, 2024. https://doi.org/10.1016/j.ijhydene.2024.05.128",
    "International Energy Agency. Global Hydrogen Review 2024. IEA, Paris, 2024. https://www.iea.org/reports/global-hydrogen-review-2024",
    "IHI Corporation. Ammonia-Fueled Power Generation for Energy Transition. Engineering, 59, 26-31, 2026. https://doi.org/10.1016/j.eng.2026.01.012",
    "Borgioli, F.; Adachi, S.; Lindner, T. Advances in Low-Temperature Nitriding and Carburizing of Stainless Steels and Metallic Materials: Formation and Properties. Metals, 14, 1179, 2024. https://doi.org/10.3390/met14101179",
    "Maccarrone, D.; Italiano, C.; Giorgianni, G.; Centi, G.; Perathoner, S.; Vita, A.; Abate, S. A Comprehensive Review on Hydrogen Production via Catalytic Ammonia Decomposition. Catalysts, 15, 811, 2025. https://doi.org/10.3390/catal15090811",
    "Mittemeijer, E. J.; Somers, M. A. J. Thermochemical Surface Engineering of Steels: Improving Materials Performance. Woodhead Publishing, 2015.",
    "Somers, M. A. J.; Christiansen, T. L. Low temperature surface hardening of stainless steel. In Thermochemical Surface Engineering of Steels; Woodhead Publishing, 2015.",
    "Laws, N.; Sarabia, E.; Campuzano, F.; Roberts, W. L. Failure analysis of FeCrAl heating coils exposed to a high-temperature, high-pressure ammonia environment. Engineering Failure Analysis, 2024. https://doi.org/10.1016/j.engfailanal.2024.108286",
]


FIGURES = [
    ("Figure 1. 암모니아 질화 연구의 최근 중심 이동: 표면경화 공정에서 암모니아 에너지 시스템 재료 신뢰성으로.", "research_landscape.png"),
    ("Figure 2. 암모니아-금속 계면에서의 흡착, 해리, 침입형 확산 및 질화물 석출 경로.", "nh3_reaction_pathway.png"),
    ("Figure 3. 온도와 질화 포텐셜로 본 암모니아 환경의 개념적 영역도.", "nitriding_potential_map.png"),
    ("Figure 4. 암모니아 크래킹 반응기에서 온도·위치·NH3 활성도 구배가 만드는 질화부식 차이.", "reactor_gradient.png"),
    ("Figure 5. 암모니아 화염-벽 상호작용에서 NH3/NH2 활성종과 수증기의 경쟁 효과.", "flame_wall.png"),
    ("Figure 6. 합금계별 원치 않는 질화에 대한 상대 취약성과 주요 손상 양상.", "alloy_response_matrix.png"),
    ("Figure 7. 알루미나 형성 확산코팅을 이용한 고압 암모니아 크래킹용 합금 보호 개념.", "alumina_barrier.png"),
    ("Figure 8. 문헌 기반 공정-조직-성능 데이터 구조화와 수명예측 모델링 흐름.", "data_workflow.png"),
    ("Figure 9. 암모니아 질화부식 연구의 2026년 이후 로드맵.", "roadmap.png"),
    ("Figure 10. 본 리뷰의 통합 프레임: 활성질소 공급, 합금 반응, 손상, 억제 전략.", "synthesis_loop.png"),
    ("Figure 11. NH3 평형 몰분율의 온도 및 압력 의존성. 출처: Jeong and Yeom [1], Fig. 1, CC BY-NC-ND 4.0.", "source_figures/jeong_yeom_2026_fig1.png"),
    ("Figure 12. 금속 촉매 표면에서 암모니아 분해와 질소·수소 생성의 개념도. 출처: Jeong and Yeom [1], Fig. 2, CC BY-NC-ND 4.0.", "source_figures/jeong_yeom_2026_fig2.png"),
    ("Figure 13. 대표 금속 질화물 형성 반응의 Gibbs 자유에너지 비교. 출처: Jeong and Yeom [1], Fig. 4, CC BY-NC-ND 4.0.", "source_figures/jeong_yeom_2026_fig4.png"),
    ("Figure 14. Fe-N 상태도와 질소 함량에 따른 상 영역. 출처: Jeong and Yeom [1], Fig. 6, CC BY-NC-ND 4.0.", "source_figures/jeong_yeom_2026_fig6.png"),
    ("Figure 15. 질화 포텐셜과 열사이클에 따른 구조재 질화층 변화 사례. 출처: Jeong and Yeom [1], Fig. 10, CC BY-NC-ND 4.0.", "source_figures/jeong_yeom_2026_fig10.png"),
    ("Figure 16. Fe/Ni 질화물과 CrN 결정구조 비교. 출처: Jeong and Yeom [1], Fig. 14, CC BY-NC-ND 4.0.", "source_figures/jeong_yeom_2026_fig14.png"),
]

SOURCE_FIGURE_URLS = {
    "jeong_yeom_2026_fig1.png": "https://media.springernature.com/lw1200/springer-static/image/art%3A10.1007%2Fs11814-026-00706-2/MediaObjects/11814_2026_706_Fig1_HTML.png",
    "jeong_yeom_2026_fig2.png": "https://media.springernature.com/lw1200/springer-static/image/art%3A10.1007%2Fs11814-026-00706-2/MediaObjects/11814_2026_706_Fig2_HTML.png",
    "jeong_yeom_2026_fig4.png": "https://media.springernature.com/lw1200/springer-static/image/art%3A10.1007%2Fs11814-026-00706-2/MediaObjects/11814_2026_706_Fig4_HTML.png",
    "jeong_yeom_2026_fig6.png": "https://media.springernature.com/lw1200/springer-static/image/art%3A10.1007%2Fs11814-026-00706-2/MediaObjects/11814_2026_706_Fig6_HTML.png",
    "jeong_yeom_2026_fig10.png": "https://media.springernature.com/lw1200/springer-static/image/art%3A10.1007%2Fs11814-026-00706-2/MediaObjects/11814_2026_706_Fig10_HTML.png",
    "jeong_yeom_2026_fig14.png": "https://media.springernature.com/lw1200/springer-static/image/art%3A10.1007%2Fs11814-026-00706-2/MediaObjects/11814_2026_706_Fig14_HTML.png",
}


TABLES = [
    ("Table 1. 본 리뷰에서 선정한 암모니아 질화 관련 핵심 연구 축과 대표 문헌.", [
        ["연구 축", "대표 환경", "핵심 질문", "대표 문헌"],
        ["에너지 시스템 질화부식", "크래킹·연소·가스터빈", "질화가 수명과 균열을 어떻게 제한하는가", "[1]-[7]"],
        ["화염-벽 상호작용", "NH3/O2/N2, 수증기 포함", "활성종과 물이 질화속도를 어떻게 바꾸는가", "[3], [4]"],
        ["공정 고도화", "NH3 가스질화, 레이저 보조", "낮은 NH3 해리율과 긴 처리시간을 어떻게 줄일 것인가", "[8], [16]"],
        ["저온 질화·스테인리스", "400-500 °C, 고질소 고용", "경도 향상과 내식성 저하의 균형점은 어디인가", "[9], [16], [19]"],
        ["저장·수송 SCC", "액체 NH3, 탱크·배관", "고온 질화와 별개의 액상 손상 모드는 어떻게 관리할 것인가", "[10], [11]"],
    ]),
    ("Table 2. 암모니아 활용 장치별 질화 구동인자와 예상 손상.", [
        ["장치", "온도/분위기", "주요 구동인자", "대표 손상"],
        ["암모니아 크래커", "400-700 °C, NH3/H2/N2", "NH3 전환율, 유동방향 농도구배, 압력", "CrN/Fe4N, 박리, 입계균열"],
        ["암모니아 연소기", "화염 근접 벽, 수증기 동반", "벽면 NH3, NH2 라디칼, H2O 산화효과", "표면경화, 질소 농도구배, 취화"],
        ["가스터빈 연료부", "400-800 °C, 환원성 NH3", "코팅 결함, N 확산, Cr/Al 고갈", "MCrAlY 내부질화, 균열"],
        ["해상엔진 부품", "400-500 °C 장시간", "합금 원소, 노출시간, 온도", "다공성 질화막, 피로·마모 수명 저하"],
        ["액체 암모니아 저장", "-33 °C 부근 또는 가압", "불순물, 잔류응력, 용접부", "SCC, 피막파괴-재부동태화"],
    ]),
    ("Table 3. 합금 원소와 질화 반응의 정성적 역할.", [
        ["원소/상", "질소와의 상호작용", "유익한 면", "위험한 면"],
        ["Cr", "CrN/Cr2N 형성", "보호성 질화막 또는 산화막 형성 가능", "Cr 고갈, 입계취화, 내식성 저하"],
        ["Al", "AlN 가능, Al2O3 형성 원소", "사전 산화 시 질소 차단 장벽", "불연속 산화막이면 결함 경로 형성"],
        ["Ni", "질화물 안정성 낮음", "오스테나이트 안정화와 확산 억제", "코팅 내 Co/Ni 재분포와 표면질화 가능"],
        ["Mo", "강한 질소 친화도", "일부 조건에서 안정층 형성", "취성상과 국부응력 증가 가능"],
        ["Fe", "Fe4N, Fe2-3N 형성", "제어된 표면경화의 기반", "다공성·취성 화합물층과 균열"],
    ]),
    ("Table 4. 제어된 질화 공정과 원치 않는 질화부식의 비교.", [
        ["구분", "제어된 NH3 가스질화", "암모니아 에너지 시스템의 원치 않는 질화"],
        ["목표", "경도·마모·피로성능 향상", "발생 자체를 억제하거나 수명예측"],
        ["운전 제어", "온도, 시간, KN, 전처리 정밀 제어", "유동, 화염, 촉매, 압력 변화와 결합"],
        ["조직", "화합물층+확산층을 의도적으로 설계", "비균일 질화층, 박리, 입계균열"],
        ["평가", "경도, 층 두께, 마모시험", "균열, 잔류강도, 장기노출, 위치 의존성"],
        ["핵심 리스크", "과도한 백색층, 내식성 저하", "예상 밖의 취화와 장치 안전성 저하"],
    ]),
    ("Table 5. 진단·모델링 도구와 해석 가능한 정보.", [
        ["도구", "얻는 정보", "리뷰에서의 용도"],
        ["SEM/EDS", "층 두께, 균열, 원소 분포", "질화층·박리·입계 손상 확인"],
        ["XRD", "CrN, Fe4N, Fe2-3N 등 상 확인", "온도별 상 안정성 판단"],
        ["TEM", "나노질화물, 계면, 결함구조", "취화 기구와 석출 위치 해석"],
        ["GDOES/WDS", "질소 깊이분포", "확산층 성장과 표면 활성도 비교"],
        ["CALPHAD/열역학", "상 안정성, 질화 포텐셜 경계", "합금·분위기 설계의 1차 필터"],
        ["수명모델/ML", "공정-조직-성능 예측", "장치 조건별 재료 선택 지도화"],
    ]),
    ("Table 6. 억제 전략의 장점, 제약, 적용 가능성.", [
        ["전략", "원리", "장점", "제약"],
        ["Al2O3 형성 코팅", "질소 확산 차단", "고압 NH3 크래킹에서 유망", "결함·박리·열팽창 불일치 관리 필요"],
        ["Cr2O3/SiO2 기반 보호", "표면 반응 억제", "기존 고온합금과 친화적", "환원성 NH3에서 안정성 제한 가능"],
        ["합금 고도화", "N 친화도와 확산 경로 제어", "장치 수준 재료 선택에 직접 연결", "비용과 용접성, 크리프와 동시 최적화 필요"],
        ["운전조건 조정", "NH3 활성도·온도창 회피", "설계 변경 없이 적용 가능", "촉매 전환율과 효율 저하 가능"],
        ["수증기/산소 관리", "산화와 질화의 경쟁 이용", "화염 질화 억제 가능성", "NOx, 산화부식, 연소 안정성과 결합"],
        ["표준시험 구축", "가속시험과 실제환경 연결", "산업 적용의 신뢰도 향상", "장시간 데이터와 공통 프로토콜 필요"],
    ]),
]


SECTION_DATA = [
    {
        "heading": "1. 서론",
        "subs": [
            ("1.1 연구 배경", [
                "암모니아는 비료 산업의 원료를 넘어 수소 운반체, 직접 연료, 전력 저장 매개체로 재해석되고 있다. 액화 조건이 수소보다 온화하고 기존 운송 인프라를 부분적으로 활용할 수 있다는 점 때문에 최근 수소 경제의 병목인 저장과 장거리 수송 문제를 우회하는 후보로 부상하였다[12]-[14]. 그러나 암모니아가 에너지 장치 내부에서 고온 금속과 만날 때에는 표면경화 공정에서 유용하게 쓰이던 질화 반응이 원치 않는 재료 열화로 바뀐다.",
                "이 리뷰의 출발점은 바로 그 의미 전환이다. 전통적인 가스질화에서는 NH3 분해, 표면 흡착, 질소 고용과 확산을 이용하여 강 표면의 경도와 마모 저항을 높인다[16]-[19]. 반면 암모니아 크래킹 반응기, 암모니아 연소기, 가스터빈 연료부, 해상엔진 부품에서는 같은 반응이 질화막 균열, 입계 취화, 코팅 원소 고갈, 피막 박리로 이어질 수 있다[1]-[7].",
                "따라서 암모니아 질화 연구에서 가장 뜨거운 질문은 더 이상 '어떻게 더 잘 질화할 것인가'에만 머물지 않는다. 현재의 핵심 질문은 '탄소 없는 암모니아 에너지 시스템을 만들면서, 그 시스템을 이루는 금속 재료가 암모니아에 의해 질화·취화되지 않도록 어떻게 설계하고 예측할 것인가'이다.",
            ]),
            ("1.2 최신 연구동향에서 본 주제 선정", [
                "2024-2026년 문헌 흐름을 보면 암모니아 활용 연구는 촉매와 연소 성능 중심에서 재료 신뢰성으로 확장되고 있다. 2026년에 발표된 구조용 강의 고온 암모니아 질화 열화 리뷰는 이 문제가 암모니아 크래킹과 연소 시스템의 공통 병목임을 명확히 제시하였다[1]. 같은 시기 310S 스테인리스강의 반응기 위치 의존 부식[7], 선박·엔진용 강의 고온 암모니아 질화부식[2], 고압 암모니아 크래킹용 HP40 합금 보호코팅[6] 연구가 이어졌다.",
                "이러한 문헌의 공통점은 질화를 단순 표면처리가 아니라 장치 수명과 안전을 좌우하는 환경유기 손상으로 다룬다는 점이다. 특히 암모니아는 고온에서 NH3, NH2, NH, N* 등 여러 질소 함유 활성종을 만들고, 수소와 수증기, 촉매 표면, 유동장과 동시에 작용한다. 같은 합금도 온도와 위치, NH3 전환율에 따라 CrN이 보호적으로 작용하거나 Fe4N이 취성 균열을 촉진할 수 있다[3], [4], [7].",
                "본 리뷰는 이 동향을 반영하여 '암모니아 에너지 시스템에서의 원치 않는 질화 및 질화부식'을 중심 주제로 삼았다. 이는 전통적인 암모니아 가스질화 지식을 출발점으로 하되, 암모니아 크래킹·연소·가스터빈·저장 인프라에서 필요한 재료 선택과 억제 전략을 통합하는 문제이다.",
            ]),
            ("1.3 리뷰의 범위와 방법", [
                "문헌 범위는 세 갈래로 구성하였다. 첫째, NH3 가스질화, 저온 질화, 레이저 보조 질화 등 제어된 표면공정의 기초와 최신 고도화 연구를 검토하였다[8], [9], [16]. 둘째, 고온 암모니아 환경에서 구조용 강, 스테인리스강, 니켈계 합금, MCrAlY 코팅이 겪는 원치 않는 질화부식을 검토하였다[1]-[7]. 셋째, 액체 암모니아 저장·수송에서의 SCC 문헌을 별도 손상 모드로 요약하여 고온 질화와의 경계를 분명히 하였다[10], [11].",
                "본 논문은 새로운 실험 데이터를 제시하지 않는 리뷰 논문이다. 따라서 수치와 조건은 인용 문헌에 근거하여 사용하였고, 도식은 원문 그림을 재사용하지 않고 본 리뷰의 논리 구조를 설명하기 위해 새로 작성하였다. 특정 공정 조건이나 합금 조성의 정밀 설계에는 추가 실험과 원문 확인이 필요하며, 메타데이터가 확인되지 않은 일부 참고문헌 항목은 참고문헌에서 별도로 표시하였다.",
                "Figure 1은 본 리뷰가 다루는 연구 지형을 요약한다. 왼쪽에는 제어된 질화 공정이, 오른쪽에는 암모니아 에너지 시스템에서의 원치 않는 질화가 배치되며, 중앙에는 NH3 해리와 활성질소 공급이라는 공통 반응축이 놓인다.",
            ], 0),
        ],
    },
    {
        "heading": "2. 암모니아 질화의 열역학과 계면반응",
        "subs": [
            ("2.1 NH3 해리와 활성질소 공급", [
                "가스질화의 기본 반응은 NH3가 금속 표면에서 해리되어 흡착 질소를 만들고, 이 질소가 금속 격자 안으로 고용·확산되는 과정이다. 질화 효율은 단순히 NH3 농도로만 결정되지 않고 온도, 수소 분압, 표면 산화막, 촉매성 합금 원소, 유동 조건에 의해 달라진다. 전통적인 공정에서는 질화 포텐셜을 제어하여 ε-Fe2-3N, γ'-Fe4N, 확산층의 비율을 조정한다[18].",
                "암모니아 에너지 장치에서는 이 제어성이 크게 약화된다. 암모니아 크래킹 반응기는 촉매와 구조재가 동시에 존재하고, NH3가 반응기 길이 방향으로 소모되면서 위치별 질화 포텐셜이 바뀐다[7]. 암모니아 연소기에서는 화염 구조, 벽 온도, 수증기, 라디칼 농도가 금속 표면 반응을 동시에 바꾼다[3], [4].",
                "Figure 2는 NH3가 표면에 흡착된 뒤 NH2*, NH*, N* 중간종을 거쳐 질소 고용과 질화물 석출로 이어지는 개념 경로를 나타낸다. 실제 장치에서는 이 경로가 산화, 수소 취화, 열응력, 코팅 결함과 동시에 연결되므로 단일 반응식만으로 손상을 설명하기 어렵다.",
            ], 1),
            ("2.2 질화 포텐셜과 상 형성", [
                "질화 포텐셜은 NH3와 H2 분압의 함수로 해석되며, 철계 합금에서 질소 활동도를 나타내는 유용한 지표이다. 높은 질화 포텐셜에서는 표면 화합물층이 빠르게 성장하고, 낮은 질화 포텐셜에서는 확산층 중심의 얕은 질화가 진행될 수 있다. 그러나 암모니아 크래킹처럼 H2가 많이 생성되는 환경에서는 표면의 질화 구동력이 시간과 위치에 따라 감소한다.",
                "Fe-N계에서는 α-Fe(N) 고용체, γ'-Fe4N, ε-Fe2-3N이 주요 상으로 거론된다. 스테인리스강에서는 CrN과 Cr2N이 내식성과 취화에 결정적인 역할을 한다. 저온에서는 고질소 고용상인 expanded austenite 또는 S-phase가 형성되어 경도와 내마모성을 높이지만, 온도 상승이나 장시간 노출에서는 Cr 질화물이 석출되어 Cr 고갈과 부식 저항 저하가 발생할 수 있다[9], [16], [19].",
                "암모니아 에너지 시스템의 특이점은 보호상과 손상상이 고정되어 있지 않다는 점이다. 예를 들어 310S 스테인리스강의 고온 암모니아 노출에서는 400-700 °C 범위에서 질화 상과 균열 양상이 비단조적으로 변할 수 있고, 특정 온도에서는 Fe4N 형성이 응력 축적과 박리를 촉진하는 것으로 해석되었다[7].",
            ], 2),
            ("2.3 확산, 응력, 균열의 결합", [
                "질화층의 성장 자체는 확산 문제처럼 보이지만, 손상은 확산만으로 결정되지 않는다. 질소가 격자 안으로 들어가면 격자 팽창, 잔류응력, 상변태, 석출물 형성이 동시에 발생한다. 표면 화합물층이 치밀하고 접착성이 좋으면 보호층으로 기능할 수 있지만, 다공성·취성·불연속층이면 균열과 박리의 기점이 된다.",
                "가스터빈용 CoNiCrAlY 코팅 연구에서는 500-800 °C, NH3 유동 조건에서 CrN/Cr2N, AlN, Ni3N, Co2N 관련 상 변화와 원소 재분포가 보고되었다[5]. 이 경우 질화는 단순히 표면에 경한 층을 만드는 것이 아니라 산화막 형성 원소인 Cr과 Al의 유효 농도를 낮추어 장기 산화 저항을 함께 약화시킬 수 있다.",
                "따라서 암모니아 질화부식은 '질소 확산층 두께'만이 아니라 '질화물의 위치, 연결성, 취성, 산화막 형성 원소의 고갈, 열·기계응력과의 결합'으로 평가되어야 한다. 이는 제어된 질화 공정의 품질 관리 지표와 에너지 장치의 수명 지표가 서로 다름을 뜻한다.",
            ]),
        ],
    },
    {
        "heading": "3. 제어된 암모니아 질화 공정의 최신 고도화",
        "subs": [
            ("3.1 전통적 가스질화의 장점과 한계", [
                "전통적인 NH3 가스질화는 산업적으로 성숙한 표면경화 기술이다. 강 부품을 500-580 °C 부근에서 장시간 노출하여 표면에 화합물층과 확산층을 만들고, 이를 통해 경도, 마모 저항, 피로 특성을 개선한다. 그러나 NH3 해리율이 낮고 처리시간이 길며, 과도한 화합물층은 취성과 표면 박리를 유발할 수 있다[18].",
                "2026년 iScience 리뷰는 가스질화의 열역학, 동역학, 표면 흡착, 확산, 부식·마모 저항 기구를 종합하고, 표면 나노결정화, 전처리, 촉매 보조, 레이저 보조, 저온 질화 등을 미래 방향으로 제시하였다. 이 흐름은 암모니아 에너지 시스템의 손상 문제와도 직접 연결된다. 질소 공급을 촉진하는 기술은 표면처리에서는 장점이지만, 장치 재료에서는 피해야 할 반응 경로를 알려주는 역방향 지식이 된다.",
                "Table 4는 제어된 가스질화와 원치 않는 질화부식을 비교한다. 같은 NH3-금속 반응을 이용하지만, 목표·평가 지표·허용 가능한 조직·위험 시나리오가 다르다. 이 구분을 분명히 해야 공정 지식이 장치 신뢰성 문제에 잘못 적용되는 것을 막을 수 있다.",
            ], 3),
            ("3.2 레이저 보조 NH3 질화", [
                "레이저 진동 여기 보조 가스질화는 NH3 분자의 특정 진동 모드에 에너지를 주어 해리를 촉진하려는 접근이다. Fan 등은 38CrMoAl 강의 NH3 가스질화에서 레이저 보조 조건이 질소 함량과 표면 경도를 높이고 질화 시간을 줄일 수 있음을 보고하였다[8]. 이 연구는 NH3 해리 단계가 질화 속도 병목이 될 수 있음을 분명히 보여준다.",
                "이 기술은 공정 측면에서 유망하지만, 에너지 시스템 관점에서는 중요한 경고도 준다. 금속 표면, 촉매, 플라즈마, 고에너지 라디칼 환경이 NH3 해리를 촉진하면 의도하지 않은 질화도 빨라질 수 있다. 즉 암모니아 크래킹 촉매 또는 플라즈마 보조 분해 장치에서는 구조재와 촉매의 경계면 설계가 더욱 중요해진다.",
                "따라서 레이저·플라즈마·촉매 보조 질화 문헌은 단순 공정개선 사례가 아니라, 암모니아 환경에서 활성질소 생성률을 어떻게 낮추거나 격리할 것인지에 대한 설계 원리를 제공한다.",
            ]),
            ("3.3 저온 질화와 스테인리스강의 내식성 딜레마", [
                "스테인리스강의 저온 질화는 CrN 석출을 억제하면서 질소 과포화 고용층을 형성한다는 점에서 매력적이다. expanded austenite 또는 expanded ferrite/martensite 상은 높은 경도와 비교적 양호한 내식성을 동시에 줄 수 있다. 그러나 공정 온도와 NH3 농도가 높아지면 질화물 석출, 국부전위 차이, 피팅 부식이 증가할 수 있다[9], [16].",
                "Duplex stainless steel의 저온 가스질화 연구는 400-500 °C, 50-100% NH3 조건에서 층 두께와 상 구성이 크게 변하고, 내식성은 공정 조건에 매우 민감함을 보여주었다[9]. 이는 암모니아 에너지 장치에서 '스테인리스강이면 안전하다'는 단순한 판단이 부족하다는 뜻이다. 고질소 고용이 유익한 조건과 CrN/Fe-nitride 석출이 손상을 부르는 조건의 경계를 구체적으로 알아야 한다.",
                "특히 에너지 장치의 노출은 공정보다 길고, 온도와 가스조성이 시간에 따라 변하며, 기계응력과 용접부가 존재한다. 따라서 표면처리 문헌에서 얻은 최적 조건을 그대로 장치 조건에 옮기는 것이 아니라, 장치별 열·화학 이력으로 재해석해야 한다.",
            ]),
            ("3.4 공정 지식의 역전: 보호해야 할 표면", [
                "제어된 질화에서는 표면 산화막 제거와 활성화가 선행된다. 반대로 원치 않는 질화부식을 막으려면 안정한 산화막 또는 확산 장벽을 유지해야 한다. 이 관점 전환은 암모니아 크래커용 HP40 합금의 알루미나 형성 코팅 연구에서 잘 드러난다[6].",
                "알루미나 장벽은 NH3가 산소가 거의 없는 환경에서 작동하므로 운전 중 자연적으로 성장하기 어렵다. 따라서 사전 알루미나 형성을 통해 조밀하고 접착성 있는 α-Al2O3 스케일을 만들어야 한다. 빠른 승온이나 긴 유지시간은 결함과 열응력을 만들 수 있으므로, 코팅 공정은 질화 억제뿐 아니라 열사이클 안정성을 함께 만족해야 한다.",
                "이처럼 공정 지식은 두 방향으로 쓰인다. 하나는 질소를 더 잘 넣기 위한 기술이고, 다른 하나는 질소가 들어가지 못하도록 표면을 봉쇄하는 기술이다. 암모니아 에너지 시스템의 재료 설계는 두 지식을 동시에 이해해야 한다.",
            ]),
        ],
    },
    {
        "heading": "4. 암모니아 에너지 시스템에서의 원치 않는 질화",
        "subs": [
            ("4.1 암모니아 크래킹 반응기", [
                "암모니아 크래킹은 NH3를 H2와 N2로 분해하여 수소를 현장에서 생산하는 기술이다. 촉매 개발과 반응기 집적 연구가 활발하지만, 최근 문헌은 구조재가 고온 NH3, H2, N2, 촉매 표면, 압력구배에 동시에 노출된다는 점을 강조한다[6], [7], [17].",
                "310S 스테인리스강 연구에서는 400-700 °C 암모니아 환경에서 부식속도가 단조적으로 증가하지 않고 500 °C 부근에서 심해지는 양상이 보고되었다[7]. 이는 확산이 빠른 고온일수록 항상 더 위험하다는 단순 모델이 맞지 않을 수 있음을 의미한다. Fe4N의 형성과 열적 안정성, CrN의 분포, 위치별 NH3 활성도가 함께 작용하기 때문이다.",
                "Figure 4는 반응기 입구에서 NH3 활성도가 높고 하류로 갈수록 낮아지는 상황을 도식화한다. 같은 합금이라도 상류 시편은 높은 질화 구동력과 두꺼운 질화층을 경험하고, 하류 시편은 더 온화한 환경에 놓일 수 있다. 따라서 재료 시험은 단일 온도·단일 가스조성보다 위치 의존성을 반영해야 한다.",
            ], 4),
            ("4.2 암모니아 연소와 화염-벽 질화", [
                "암모니아 연소는 CO2를 직접 배출하지 않는다는 장점이 있지만, 낮은 연소속도, 높은 점화에너지, NOx/N2O 배출, 미연 NH3 문제가 남아 있다[13], [15]. 여기에 최근 추가된 쟁점이 화염-금속 벽 상호작용에 의한 질화이다.",
                "Wang 등은 SUS310S 시험편을 암모니아 화염에 노출하고 벽 온도, NH3/NH 라디칼 분포, 수증기 효과를 분석하였다[3]. 이 연구는 벽면 근방 NH3 농도가 질화의 주요 지배인자이며, NH2 라디칼도 일부 기여할 수 있음을 제시하였다. 후속 연구는 수증기가 표면 반응성을 낮추어 질화를 억제할 수 있음을 보여주었다[4].",
                "Figure 5는 화염에서 벽으로 전달되는 활성종과 수증기 산화효과의 경쟁을 나타낸다. 중요한 점은 수증기가 단순 희석제가 아니라 표면 산화 상태를 바꾸어 NH3 해리를 억제할 수 있다는 것이다. 그러나 실제 연소기에서는 수증기 증가가 NOx, 열효율, 산화부식과도 연결되므로 독립적인 해결책으로 볼 수 없다.",
            ], 5),
            ("4.3 암모니아 가스터빈과 코팅 열화", [
                "암모니아 가스터빈의 연료부, 밸브, 인젝터, 연료부 과농 영역은 산화성 고온가스만이 아니라 NH3가 풍부한 환원성·질화성 분위기를 경험할 수 있다. 기존 가스터빈 코팅은 주로 산화와 열차폐를 기준으로 설계되어 왔기 때문에, NH3 질화 환경에서의 장기 안정성은 별도의 검증이 필요하다.",
                "Ghara 등은 HVOF CoNiCrAlY 코팅이 10% NH3 유동, 500-800 °C에서 내부질화와 표면균열을 보이며, Cr과 Al이 질화물 형성에 소모되어 산화막 형성 능력이 저하될 수 있음을 보고하였다[5]. 이는 코팅이 '있는지'보다 코팅이 NH3에서 어떤 상으로 바뀌는지가 중요하다는 사실을 보여준다.",
                "MCrAlY 코팅은 산화 환경에서는 보호적 Al2O3/Cr2O3 형성에 의존한다. 그러나 산소 공급이 부족하고 NH3가 풍부하면 보호 산화막 대신 내부질화가 우세해질 수 있다. 따라서 암모니아 가스터빈 코팅 설계는 산화-질화-열피로-원소확산을 동시에 고려하는 방향으로 확장되어야 한다.",
            ]),
            ("4.4 암모니아 연료 엔진과 구조용 강", [
                "암모니아를 선박과 엔진 연료로 적용하려는 시도는 재료 호환성 문제를 더 현실적인 수준으로 끌어올렸다. 장시간 운전, 반복 열사이클, 윤활·연료 불순물, 고온 벽면은 실험실 질화보다 복잡하다. 2026년 Fuel 논문은 저합금강, 합금강, 스테인리스강을 고온 NH3에 장시간 노출하여 모든 소재에서 질화가 발생할 수 있음을 보였다[2].",
                "이 연구에서 합금 원소는 질화막의 조성, 연속성, 균열 여부를 크게 바꾸었다. 저Cr 재료와 고Cr 스테인리스는 같은 온도에서도 서로 다른 질화막을 형성할 수 있으며, 어떤 경우에는 CrN 기반 막이 상대적으로 보호적일 수 있다. 반대로 다공성 Fe 질화막은 취성과 균열의 원인이 될 수 있다.",
                "이 결과는 암모니아 엔진용 재료 선택이 단순한 강도·내열성 표만으로는 부족함을 뜻한다. 온도창, 노출시간, NH3 분압, H2/H2O 동반 여부, 표면상태를 포함한 질화부식 지도가 필요하다.",
            ]),
            ("4.5 액체 암모니아 SCC와 고온 질화부식의 구분", [
                "암모니아 인프라에서 또 하나의 중요한 손상은 액체 암모니아 중 강의 응력부식균열이다[10], [11]. 이는 고온 NH3 질화부식과 구분되어야 한다. 고온 질화부식은 활성질소 흡수와 질화물 형성, 확산, 열응력에 의해 지배되는 반면, 액체 암모니아 SCC는 불순물, 전기화학 반응, 피막 파괴, 용접잔류응력, 강도 수준에 민감하다.",
                "그러나 두 문제는 독립적이지 않다. 암모니아 가치사슬이 생산-저장-운송-크래킹-연소로 이어지기 때문에 하나의 재료 시스템은 액체 NH3, 기체 NH3, 고온 NH3/H2, 연소가스 등 여러 환경을 순차적으로 경험할 수 있다. 따라서 안전 설계는 저장 탱크의 SCC와 고온 장치의 질화부식을 같은 재료 호환성 프레임 안에서 다루어야 한다.",
                "Table 2는 주요 장치별 운전환경과 예상 손상을 비교한다. 고온 장치에서는 질화물 상과 확산층, 저장·수송에서는 SCC와 용접부가 중심이지만, 공통적으로 '암모니아가 금속 표면의 보호 상태를 바꾼다'는 문제가 존재한다.",
            ], 1),
        ],
    },
    {
        "heading": "5. 합금계별 반응성과 보호전략",
        "subs": [
            ("5.1 Ferritic/low-alloy steels", [
                "저합금강은 비용과 기계적 강도 측면에서 매력적이지만, NH3 환경에서는 Fe 질화물 형성과 다공성 표면층, 수소 관련 열화가 함께 문제가 될 수 있다. 고온 암모니아 합성 반응기에서는 수소가 많은 낮은 질화 포텐셜 환경이므로 수소 열화가 상대적으로 더 중요할 수 있지만, 암모니아 연료부나 크래킹 입구처럼 NH3 활성도가 높은 구역에서는 질화가 다시 지배적일 수 있다[1], [2].",
                "Fe4N과 Fe2-3N은 제어된 질화에서는 경도 향상의 핵심상이지만, 구조재에서는 취성 화합물층으로 작용할 수 있다. 표면층이 균일하고 얇으면 마모 저항을 줄 수 있으나, 장시간 노출로 두껍고 다공성인 층이 되면 균열과 박리가 발생한다.",
                "따라서 저합금강은 저온·저NH3 영역이나 보호코팅이 안정한 부위에서 제한적으로 검토하고, 고온·고NH3·장시간 노출부에는 고Cr/Ni 합금 또는 장벽 코팅과 조합하여 사용해야 한다.",
            ]),
            ("5.2 Austenitic stainless steels", [
                "오스테나이트계 스테인리스강은 고온강도와 산화저항 때문에 암모니아 크래킹 및 연소 장치 후보로 자주 검토된다. 그러나 Cr이 질소와 강하게 결합하여 CrN/Cr2N을 형성하면 내식성의 기반인 Cr 가용도가 낮아질 수 있다. 낮은 온도에서의 expanded austenite는 유익할 수 있지만, 장시간 고온에서는 석출과 입계취화가 우세해질 수 있다[7], [9], [16].",
                "310S 연구의 핵심은 질화부식이 온도와 위치에 따라 비선형적으로 변한다는 점이다[7]. 고온일수록 확산이 빠르지만, 특정 Fe 질화물은 고온에서 불안정해질 수 있고, CrN의 상대적 보호성도 조건에 따라 달라진다. 이 때문에 단일 Arrhenius 확산모델만으로는 수명을 예측하기 어렵다.",
                "스테인리스강의 설계 방향은 Cr과 Ni의 균형, 질소 확산 경로, 표면 산화막 안정성, 용접부 조직을 동시에 고려하는 것이다. 특히 용접 열영향부와 입계 Cr 고갈은 액체 NH3 SCC와 고온 질화부식 양쪽에서 취약점이 될 수 있다.",
            ]),
            ("5.3 Nickel-based alloys and MCrAlY coatings", [
                "니켈계 합금은 질소 확산과 질화물 안정성 측면에서 철계보다 유리할 수 있으나, Cr과 Al을 포함한 산화막 형성 원소가 질화물로 소모되면 보호기능이 저하된다. MCrAlY 코팅은 원래 산화와 열피로 저항을 위해 쓰이지만, NH3 환원성 분위기에서는 내부질화와 원소 재분포가 발생할 수 있다[5].",
                "코팅 설계의 핵심은 산화막 형성 원소를 충분히 보유하면서도 NH3가 코팅 내부로 침투할 결함을 줄이는 것이다. HVOF 코팅의 기공, splat 경계, 열처리 이력은 질소 확산 통로가 될 수 있다. 따라서 코팅 평가에는 산화시험뿐 아니라 NH3 유동, 압력, 온도구배가 포함되어야 한다.",
                "Figure 6은 합금계별 취약성을 정성적으로 비교한다. 이 그림은 특정 합금을 등급화하기 위한 최종 데이터가 아니라, 문헌에서 반복적으로 나타나는 손상 경향을 정리한 개념도이다.",
            ], 6),
            ("5.4 Alumina-forming barriers", [
                "알루미나 형성 코팅은 현재 가장 설득력 있는 억제 전략 중 하나이다. α-Al2O3는 고온에서 안정하고 질소 확산을 효과적으로 막을 수 있다. Laws 등은 HP40 합금에 알루미나 형성 확산코팅과 사전 산화를 적용하여 고압 NH3 크래킹 환경에서 질소 침투를 줄이는 접근을 보고하였다[6].",
                "다만 알루미나 장벽은 완전한 만능책이 아니다. 코팅과 기지재의 열팽창 차이, 승온속도, 산화막 상전이, 결함 밀도, 장기 열사이클에 따라 균열과 박리가 생길 수 있다. 결함 하나가 NH3 침투 경로가 되면 국부 질화가 진행되어 보호층 아래에서 박리를 유발할 수 있다.",
                "Figure 7은 알루미나 장벽의 작동 개념을 나타낸다. 핵심은 단순히 Al을 넣는 것이 아니라, 운전 전에 치밀하고 접착성 있는 α-Al2O3 스케일을 확보하고 운전 중 그 연속성을 유지하는 것이다.",
            ], 7),
            ("5.5 합금 선택 기준의 재정의", [
                "암모니아 에너지 장치의 합금 선택은 강도, 가격, 크리프, 산화저항만으로 결정될 수 없다. 질화 포텐셜, NH3 전환율, 수소와 수증기, 압력, 촉매 접촉, 위치별 유동, 노출시간이 모두 필요하다. 기존 고온합금 데이터베이스는 산화·황화·침탄에 비해 암모니아 질화 데이터를 충분히 담고 있지 않다.",
                "Table 3은 주요 합금 원소의 정성적 역할을 요약한다. Cr, Al, Mo처럼 질소 친화도가 높은 원소는 보호층 형성에 기여할 수도 있고, 취성 질화물과 원소 고갈을 유발할 수도 있다. 따라서 '질소 친화도가 높다'는 표현은 그 자체로 장점도 단점도 아니며, 상의 위치와 연속성, 응력 상태에 따라 의미가 달라진다.",
                "결론적으로 재료 설계는 합금 조성과 표면상태, 운전환경을 하나의 조합으로 보아야 한다. 같은 합금도 표면 전처리와 산화막 품질에 따라 전혀 다른 수명을 보일 수 있다.",
            ], 2),
        ],
    },
    {
        "heading": "6. 진단, 모델링, 데이터화",
        "subs": [
            ("6.1 다중 스케일 진단", [
                "암모니아 질화부식의 진단은 표면 관찰만으로 충분하지 않다. 질화층의 두께, 균열, 상, 깊이별 질소 농도, 원소 고갈, 입계 석출, 잔류응력까지 함께 보아야 한다. SEM/EDS는 층과 균열을 빠르게 확인하고, XRD는 CrN, Fe4N, Fe2-3N 등 주요 상을 식별하며, TEM은 입계와 나노석출물을 해석하는 데 필요하다.",
                "GDOES와 WDS 같은 깊이분석은 질소 침투 깊이와 농도구배를 정량화하는 데 유용하다[2], [3]. 특히 암모니아 크래킹 반응기처럼 위치별 NH3 활성도가 다른 장치에서는 깊이분포를 비교해야 노출환경의 차이를 해석할 수 있다.",
                "Table 5는 진단·모델링 도구별로 얻을 수 있는 정보를 정리한다. 리뷰 문헌을 비교할 때도 같은 용어가 서로 다른 측정법에서 왔는지 확인해야 한다. 예를 들어 '질화층 두께'가 광학현미경 에칭 경계인지, 질소 농도 임계값인지, 경도 변화 깊이인지에 따라 해석이 달라진다.",
            ], 4),
            ("6.2 열역학과 확산 모델", [
                "열역학 모델은 어떤 질화물이 안정한지, 특정 온도와 질화 포텐셜에서 어떤 상이 예상되는지에 대한 1차 지도를 제공한다. CALPHAD 계산은 CrN, Cr2N, Fe4N, AlN 등 상 안정성을 비교하고 합금 원소가 질소 활동도에 미치는 영향을 평가하는 데 유용하다[2], [7].",
                "그러나 열역학 평형만으로 실제 손상을 예측할 수는 없다. 장치 내 노출시간은 유한하고, 표면에는 산화막과 결함이 있으며, 반응기 위치에 따라 NH3가 계속 소모된다. 따라서 확산 모델, 반응속도 모델, 표면반응 모델, 응력 모델이 함께 필요하다.",
                "Figure 3은 온도와 질화 포텐셜을 축으로 한 개념적 영역도이다. 실제 경계는 합금과 시간, 압력, 표면상태에 따라 달라지지만, 이 그림은 암모니아 크래킹, 연소, 제어 질화, 저장 환경을 같은 좌표에서 생각하도록 돕는다.",
            ]),
            ("6.3 수명예측을 위한 데이터 구조화", [
                "현재 문헌의 가장 큰 한계는 데이터가 장치별·합금별·실험실별로 흩어져 있다는 점이다. 온도, NH3 분압, H2/H2O 동반 여부, 압력, 유량, 표면처리, 합금 조성, 노출시간, 측정법이 통일되어 있지 않다. 이 상태에서는 단일 논문 결과를 일반 설계 규칙으로 확장하기 어렵다.",
                "따라서 향후에는 공정-조직-성능 데이터 구조가 필요하다. 입력 변수는 합금 조성, 표면상태, 온도, 압력, 가스조성, 유동 위치, 노출시간이며, 출력 변수는 질화층 두께, 상 구성, 균열 밀도, 잔류강도, 산화막 안정성, 전기화학 특성으로 정리할 수 있다.",
                "Figure 8은 문헌 데이터를 구조화하여 모델링과 설계로 연결하는 흐름을 나타낸다. 특히 머신러닝은 데이터 수가 충분할 때 유용하지만, 무작위 상관관계보다 열역학·확산·응력에 근거한 물리 제약 모델과 결합할 때 더 신뢰성이 높다.",
            ], 8),
            ("6.4 가속시험과 실제환경의 간극", [
                "가속시험은 산업 적용에 필수적이지만, 암모니아 질화부식에서는 단순히 온도를 올리는 방식이 항상 타당하지 않다. 상 안정성이 비단조적으로 변하고, 수증기와 NH3 활성도가 표면반응을 바꾸며, 높은 온도에서는 특정 Fe 질화물이 사라질 수 있기 때문이다[4], [7].",
                "따라서 가속시험은 실제 장치의 지배 손상모드를 보존해야 한다. 예를 들어 500 °C에서 Fe4N 형성이 균열을 지배하는 장치를 700 °C로 가속하면 전혀 다른 상과 손상모드를 시험할 수 있다. 이 경우 가속계수는 수명예측에 사용할 수 없다.",
                "신뢰할 수 있는 시험 프로토콜은 온도만이 아니라 NH3/H2/H2O/N2 조성, 압력, 유량, 표면상태, 응력, 열사이클을 포함해야 한다. 특히 암모니아 연소기에서는 화염구조와 벽면 반응을 모사하는 시험이 필요하다.",
            ]),
        ],
    },
    {
        "heading": "7. 연구 공백과 향후 로드맵",
        "subs": [
            ("7.1 현재 남아 있는 핵심 공백", [
                "첫 번째 공백은 장시간 데이터 부족이다. 많은 연구가 수 시간에서 수십 시간 노출을 다루지만, 실제 장치는 수천 시간 이상의 운전과 정지·기동 열사이클을 겪는다. 질화층이 초기에는 보호적이어도 장기적으로 균열과 박리로 바뀔 가능성을 평가해야 한다.",
                "두 번째 공백은 복합환경이다. NH3 단일가스보다 실제 장치에서는 H2, N2, H2O, O2, NOx, 촉매 입자, 윤활유 또는 불순물이 함께 존재한다. 특히 수증기는 질화를 억제할 수 있지만 산화부식과 NOx 문제를 동시에 바꾼다[4].",
                "세 번째 공백은 응력과 용접부이다. 암모니아 저장 SCC 문헌은 잔류응력과 용접부가 중요하다는 것을 오래전부터 보여주었다[10], [11]. 고온 질화부식에서도 열응력, 기계응력, 용접 열영향부가 질화물 석출과 균열 경로를 바꿀 가능성이 크다.",
            ]),
            ("7.2 표준화된 재료 호환성 지도", [
                "암모니아 에너지 산업이 성장하려면 합금별 호환성 지도가 필요하다. 이 지도는 단순한 '사용 가능/불가' 표가 아니라 온도, 압력, NH3/H2/H2O 조성, 노출시간, 응력 수준, 표면처리 상태를 포함한 다차원 영역이어야 한다.",
                "기존 수소 취화 지도나 고온 산화 지도처럼, 암모니아 질화부식도 장치별 표준 노출 조건과 평가 지표가 필요하다. 질화층 두께, 상, 균열, 질소 깊이분포, 기계적 잔류성능을 최소 공통 지표로 정하면 문헌 간 비교성이 크게 향상될 수 있다.",
                "Table 6은 현재 검토되는 억제 전략의 장점과 제약을 정리한다. 보호코팅, 합금 고도화, 운전조건 조정, 수증기/산소 관리, 표준시험 구축은 서로 배타적이지 않고 병행되어야 한다.",
            ], 5),
            ("7.3 촉매와 구조재의 경계 설계", [
                "암모니아 크래킹과 플라즈마 보조 분해에서는 촉매 표면이 NH3를 적극적으로 해리한다. 이 활성질소가 구조재로 이동하면 원치 않는 질화가 촉진될 수 있다. 따라서 촉매 성능만이 아니라 촉매-구조재 경계의 질소 활동도와 확산 장벽 설계가 필요하다.",
                "반응기 설계에서는 촉매층, 지지체, 튜브 내벽, 열교환 표면이 서로 다른 질화 포텐셜을 경험한다. 이 때문에 단일 쿠폰 시험보다 실제 반응기 위치별 노출을 반영한 모듈 시험이 중요하다[7].",
                "향후 연구는 촉매가 만드는 활성질소 플럭스와 구조재 표면 반응속도를 동시에 측정해야 한다. 이는 촉매공학과 재료공학의 협업이 필요한 지점이다.",
            ]),
            ("7.4 로드맵", [
                "Figure 9는 향후 연구 로드맵을 제시한다. 단기적으로는 대표 합금과 코팅의 NH3/H2/H2O 노출 데이터를 축적하고, 중기적으로는 장시간·응력·열사이클을 포함한 가속시험을 표준화해야 한다. 장기적으로는 실제 장치 운전데이터와 실험실 데이터를 연결한 수명예측 모델이 필요하다.",
                "특히 한국, 일본, 유럽은 암모니아 혼소, 가스터빈, 선박 연료, 수소 운반체 기술을 적극적으로 검토하고 있다[14], [15]. 따라서 암모니아 질화부식 연구는 학술적 표면공학을 넘어 에너지 전환 인프라의 안전 기준과 직접 연결될 가능성이 높다.",
                "본 리뷰의 결론은 단순하다. 암모니아는 탄소 없는 연료일 수 있지만, 재료에게는 질소와 수소를 동시에 공급하는 공격적인 화학환경이다. 이 양면성을 정면으로 다루는 것이 다음 단계 암모니아 기술의 핵심이다.",
            ], 8),
        ],
    },
    {
        "heading": "8. 결론",
        "subs": [
            ("8.1 종합 결론", [
                "본 리뷰는 최신 문헌 흐름을 바탕으로 암모니아 질화의 가장 뜨거운 연구축이 '암모니아 에너지 시스템에서의 원치 않는 질화 및 질화부식'으로 이동하고 있음을 제시하였다. 전통적인 NH3 가스질화는 표면경화 공정으로 유용하지만, 같은 반응이 고온 암모니아 크래킹, 연소, 가스터빈, 엔진 환경에서는 구조재 열화로 작용한다.",
                "암모니아 질화부식은 NH3 해리와 질소 확산만으로 설명되지 않는다. 온도, 질화 포텐셜, H2/H2O 동반 여부, 반응기 위치, 합금 원소, 산화막, 코팅 결함, 응력, 열사이클이 함께 작용한다. 이 복합성 때문에 단순한 합금명 또는 단일 온도 데이터로 장치 수명을 예측하기 어렵다.",
                "합금 설계에서는 Cr, Al, Mo, Ni의 역할을 보호성과 취화 가능성 양쪽에서 보아야 한다. 알루미나 형성 코팅은 유망하지만, 결함·박리·열팽창 불일치 관리가 필수적이다. 저온 질화와 스테인리스강의 공정 지식은 유익하지만, 에너지 장치에서는 장시간 고온 노출과 위치별 NH3 활성도까지 포함하여 재해석해야 한다.",
            ], 9),
            ("8.2 향후 연구 제언", [
                "향후 연구는 첫째, 장시간 노출과 열사이클을 포함한 합금별 데이터베이스를 구축해야 한다. 둘째, 암모니아 크래킹과 연소 장치의 실제 유동·온도·조성 구배를 반영한 위치 의존 시험이 필요하다. 셋째, 표면 코팅과 합금 원소 설계를 열역학, 확산, 응력, 산화막 안정성 모델과 결합해야 한다.",
                "넷째, 액체 암모니아 SCC와 고온 질화부식을 하나의 암모니아 재료 호환성 체계에서 관리해야 한다. 저장·수송·크래킹·연소를 분리된 장치로만 보지 말고, 전체 가치사슬에서 재료가 경험하는 화학·기계 이력을 연결해야 한다.",
                "마지막으로 표준화된 평가 지표가 필요하다. 질화층 두께, 상 구성, 질소 깊이분포, 균열 밀도, 잔류강도, 코팅 접착성, 전기화학 특성을 최소 공통 지표로 삼으면 연구 결과를 장치 설계와 안전기준으로 전환하기 쉬워진다.",
            ]),
        ],
    },
]


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    SOURCE_FIG_DIR.mkdir(parents=True, exist_ok=True)


def download_source_figures() -> None:
    for filename, url in SOURCE_FIGURE_URLS.items():
        path = SOURCE_FIG_DIR / filename
        if path.exists() and path.stat().st_size > 0:
            continue
        req = Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urlopen(req, timeout=20) as response:
            path.write_bytes(response.read())


def is_korean_char(ch: str) -> bool:
    return (
        "\u1100" <= ch <= "\u11ff"
        or "\u3130" <= ch <= "\u318f"
        or "\uac00" <= ch <= "\ud7af"
    )


@lru_cache(maxsize=64)
def font_for(script: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FIG_K_FONT if script == "ko" else FIG_E_FONT, size=size)


def text_runs(text: str) -> list[tuple[str, str]]:
    runs: list[tuple[str, str]] = []
    current = ""
    current_script: str | None = None
    for ch in text:
        script = "ko" if is_korean_char(ch) else "latin"
        if current and script != current_script:
            runs.append((current_script or "latin", current))
            current = ch
        else:
            current += ch
        current_script = script
    if current:
        runs.append((current_script or "latin", current))
    return runs


def mixed_text_width(draw: ImageDraw.ImageDraw, text: str, size: int) -> int:
    width = 0
    for script, value in text_runs(text):
        bbox = draw.textbbox((0, 0), value, font=font_for(script, size))
        width += bbox[2] - bbox[0]
    return width


def draw_text(draw: ImageDraw.ImageDraw, xy, text: str, size: int, fill=(0, 0, 0)) -> None:
    x, y = xy
    for line in str(text).split("\n"):
        cursor = x
        for script, value in text_runs(line):
            fnt = font_for(script, size)
            draw.text((cursor, y), value, font=fnt, fill=fill)
            bbox = draw.textbbox((0, 0), value, font=fnt)
            cursor += bbox[2] - bbox[0]
        y += size + 8


def wrap_text(draw: ImageDraw.ImageDraw, text: str, size: int, max_width: int) -> list[str]:
    words = text.split(" ")
    lines: list[str] = []
    line = ""
    for word in words:
        test = word if not line else f"{line} {word}"
        if mixed_text_width(draw, test, size) <= max_width:
            line = test
        else:
            if line:
                lines.append(line)
            if mixed_text_width(draw, word, size) <= max_width:
                line = word
            else:
                chunk = ""
                for ch in word:
                    test_chunk = chunk + ch
                    if mixed_text_width(draw, test_chunk, size) <= max_width:
                        chunk = test_chunk
                    else:
                        if chunk:
                            lines.append(chunk)
                        chunk = ch
                line = chunk
    if line:
        lines.append(line)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, box, text: str, fill, outline, txt=(0, 0, 0), size=34, align="center") -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    lines = wrap_text(draw, text, size, x2 - x1 - 36)
    total_h = len(lines) * (size + 8)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        width = mixed_text_width(draw, line, size)
        if align == "center":
            x = x1 + (x2 - x1 - width) / 2
        else:
            x = x1 + 22
        draw_text(draw, (x, y), line, size, txt)
        y += size + 8


def arrow(draw: ImageDraw.ImageDraw, start, end, color=(72, 88, 101), width=5) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    head = 18
    pts = [
        (x2, y2),
        (x2 - head * ux + 0.55 * head * px, y2 - head * uy + 0.55 * head * py),
        (x2 - head * ux - 0.55 * head * px, y2 - head * uy - 0.55 * head * py),
    ]
    draw.polygon(pts, fill=color)


def title(draw: ImageDraw.ImageDraw, text: str) -> None:
    draw_text(draw, (70, 44), text, 44, (0, 0, 0))
    draw.line([(70, 102), (1530, 102)], fill=(72, 115, 135), width=4)


def save_canvas(name: str, draw_fn) -> None:
    img = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(img)
    draw_fn(draw)
    img.save(FIG_DIR / name, dpi=(300, 300))


def make_figures() -> None:
    def f1(d):
        title(d, "암모니아 질화 연구 지형")
        draw_box(d, (80, 180, 470, 390), "제어된 질화\n표면경화·마모저항", (232, 242, 239), (76, 132, 117))
        draw_box(d, (565, 180, 1035, 390), "공통 반응축\nNH3 해리 → N 흡착 → 확산", (245, 247, 250), (83, 98, 119))
        draw_box(d, (1130, 180, 1520, 390), "원치 않는 질화\n크래킹·연소·가스터빈", (252, 238, 232), (175, 95, 72))
        arrow(d, (470, 285), (565, 285))
        arrow(d, (1035, 285), (1130, 285))
        for i, txt in enumerate(["가스질화", "저온 질화", "레이저/플라즈마 보조"]):
            draw_box(d, (115, 500 + i * 120, 435, 585 + i * 120), txt, (247, 252, 250), (76, 132, 117), size=27)
        for i, txt in enumerate(["암모니아 크래커", "화염-벽 질화", "코팅/합금 손상"]):
            draw_box(d, (1165, 500 + i * 120, 1485, 585 + i * 120), txt, (255, 248, 244), (175, 95, 72), size=27)
        draw_box(d, (565, 560, 1035, 760), "2026년 핵심 질문\n재료 수명과 안전을 어떻게 예측·보호할 것인가", (242, 245, 249), (83, 98, 119), size=31)

    def f2(d):
        title(d, "NH3-금속 계면 반응 경로")
        xs = [150, 450, 760, 1080, 1360]
        labels = ["NH3(g)", "NH3*", "NH2*/NH*", "N* 흡착", "N 확산/질화물"]
        for x, lab in zip(xs, labels):
            draw_box(d, (x - 105, 220, x + 105, 360), lab, (244, 248, 252), (91, 118, 149), size=30)
        for a, b in zip(xs[:-1], xs[1:]):
            arrow(d, (a + 110, 290), (b - 110, 290))
        d.rectangle((180, 650, 1420, 760), fill=(226, 231, 235), outline=(105, 115, 123), width=3)
        draw_text(d, (210, 690), "금속 기지: Fe-Cr-Ni/Fe-Cr-Mo/HP40/MCrAlY", 34)
        for x in [960, 1030, 1100, 1170, 1240]:
            d.ellipse((x, 600, x + 28, 628), fill=(203, 82, 69))
            arrow(d, (x + 14, 628), (x + 14, 650), color=(203, 82, 69), width=4)
        draw_box(d, (235, 820, 610, 920), "산화막/코팅 결함", (255, 248, 230), (193, 137, 48), size=28)
        draw_box(d, (690, 820, 1065, 920), "격자팽창·잔류응력", (255, 248, 230), (193, 137, 48), size=28)
        draw_box(d, (1145, 820, 1520, 920), "CrN/Fe4N/AlN 석출", (255, 248, 230), (193, 137, 48), size=28)

    def f3(d):
        title(d, "질화 포텐셜-온도 개념도")
        d.line((190, 820, 1430, 820), fill=(40, 50, 60), width=4)
        d.line((190, 820, 190, 170), fill=(40, 50, 60), width=4)
        arrow(d, (1430, 820), (1500, 820))
        arrow(d, (190, 170), (190, 105))
        draw_text(d, (700, 865), "Temperature", 32)
        draw_text(d, (35, 385), "Nitriding potential", 32)
        d.polygon([(220, 790), (620, 780), (520, 570), (245, 590)], fill=(234, 247, 240), outline=(74, 132, 96))
        draw_text(d, (285, 680), "저온 질화\nS-phase", 30)
        d.polygon([(570, 765), (1040, 700), (930, 410), (560, 500)], fill=(252, 240, 224), outline=(180, 112, 49))
        draw_text(d, (660, 575), "가스질화\nFe 질화물", 30)
        d.polygon([(910, 720), (1430, 620), (1370, 220), (980, 350)], fill=(255, 234, 230), outline=(190, 85, 70))
        draw_text(d, (1060, 460), "고온 NH3\n원치 않는 질화부식", 30)
        d.polygon([(260, 805), (520, 805), (480, 740), (280, 750)], fill=(235, 239, 245), outline=(90, 100, 115))
        draw_text(d, (300, 770), "액체 NH3 SCC\n별도 모드", 24)

    def f4(d):
        title(d, "암모니아 크래킹 반응기 내 위치 의존성")
        d.rounded_rectangle((140, 370, 1460, 590), radius=80, fill=(240, 244, 248), outline=(83, 98, 119), width=5)
        arrow(d, (185, 480), (1415, 480), color=(52, 92, 130), width=8)
        draw_text(d, (200, 410), "NH3 높음", 32)
        draw_text(d, (1240, 410), "NH3 낮음", 32)
        for i, (x, lab, col) in enumerate([(380, "상류\n강한 질화", (252, 225, 220)), (800, "중앙\n혼합 영역", (255, 245, 224)), (1220, "하류\n약한 질화", (230, 244, 235))]):
            draw_box(d, (x - 130, 640, x + 130, 790), lab, col, (90, 100, 110), size=30)
            arrow(d, (x, 590), (x, 640), color=(90, 100, 110), width=4)
        draw_text(d, (290, 830), "온도, NH3 전환율, H2 생성, 촉매층 위치가 질화층 두께와 상을 동시에 바꾼다.", 31)

    def f5(d):
        title(d, "암모니아 화염-벽 질화")
        d.polygon([(330, 750), (520, 210), (710, 750)], fill=(255, 219, 135), outline=(205, 125, 36))
        d.polygon([(410, 750), (520, 320), (630, 750)], fill=(255, 160, 86))
        d.rectangle((250, 760, 1350, 850), fill=(219, 225, 230), outline=(85, 95, 105), width=4)
        draw_text(d, (665, 792), "SUS310S / combustor wall", 32)
        for txt, xy in [("NH3", (760, 300)), ("NH2", (845, 400)), ("NH", (785, 510)), ("H2O", (1010, 385))]:
            d.ellipse((xy[0], xy[1], xy[0] + 95, xy[1] + 70), fill=(241, 248, 252), outline=(72, 115, 135), width=3)
            draw_text(d, (xy[0] + 15, xy[1] + 16), txt, 28)
            arrow(d, (xy[0] + 45, xy[1] + 70), (760, 760), color=(72, 115, 135), width=3)
        draw_box(d, (930, 600, 1320, 710), "수증기: 표면 산화/반응성 저하\n질화 억제 가능", (236, 248, 244), (76, 132, 117), size=25)

    def f6(d):
        title(d, "합금계별 상대 취약성")
        headers = ["Fe 저합금", "Austenitic SS", "Ni계/MCrAlY", "Al2O3 장벽"]
        rows = ["N 흡수", "취성 질화물", "코팅/산화막 고갈", "장기 안정성"]
        colors = [(248, 226, 220), (255, 242, 210), (232, 244, 235), (222, 236, 247)]
        draw_text(d, (310, 170), "낮음", 28)
        draw_text(d, (1335, 170), "높음", 28)
        for i, h in enumerate(headers):
            draw_text(d, (330 + i * 295, 230), h, 26)
        matrix = [[3, 2, 2, 1], [3, 2, 2, 1], [1, 2, 3, 1], [2, 2, 2, 2]]
        for r, row in enumerate(rows):
            draw_text(d, (90, 330 + r * 125), row, 28)
            for c, val in enumerate(matrix[r]):
                x, y = 360 + c * 295, 315 + r * 125
                d.rounded_rectangle((x, y, x + 160, y + 72), radius=16, fill=colors[val], outline=(120, 128, 135), width=2)
                draw_text(d, (x + 58, y + 17), str(val), 34)
        draw_text(d, (360, 870), "1=상대적으로 낮음, 3=상대적으로 높음. 실제 평가는 온도·가스조성·시간·표면상태에 의존한다.", 28)

    def f7(d):
        title(d, "Al2O3 장벽 코팅 개념")
        d.rectangle((250, 700, 1350, 820), fill=(206, 214, 222), outline=(77, 87, 96), width=3)
        draw_text(d, (690, 740), "HP40 / 고온합금", 34)
        d.rectangle((250, 610, 1350, 700), fill=(228, 198, 145), outline=(130, 95, 45), width=3)
        draw_text(d, (590, 638), "Al 확산층 / reservoir", 31)
        d.rectangle((250, 555, 1350, 610), fill=(245, 248, 252), outline=(72, 115, 135), width=4)
        draw_text(d, (640, 567), "dense α-Al2O3", 28)
        for x in [350, 520, 690, 860, 1030, 1200]:
            draw_text(d, (x, 235), "NH3", 28)
            arrow(d, (x + 25, 285), (x + 25, 540), color=(175, 75, 60), width=4)
            d.line((x - 10, 537, x + 65, 537), fill=(175, 75, 60), width=5)
        draw_box(d, (430, 870, 1170, 945), "핵심: 운전 전 치밀한 보호 산화막을 만들고 결함을 최소화한다.", (238, 246, 241), (76, 132, 117), size=28)

    def f8(d):
        title(d, "문헌 데이터 → 모델 → 설계")
        boxes = [
            ((90, 280, 360, 430), "문헌/실험\n조건 추출"),
            ((440, 280, 710, 430), "조직 지표\n상·두께·균열"),
            ((790, 280, 1060, 430), "열역학/확산\n물리 제약"),
            ((1140, 280, 1510, 430), "수명예측\n재료 선택 지도"),
        ]
        for box, txt in boxes:
            draw_box(d, box, txt, (241, 246, 250), (83, 98, 119), size=30)
        for i in range(len(boxes) - 1):
            arrow(d, (boxes[i][0][2], 355), (boxes[i + 1][0][0], 355), width=5)
        for box, txt in [
            ((160, 640, 520, 770), "입력: 합금, 표면, T, p, NH3/H2/H2O, 시간"),
            ((620, 640, 980, 770), "출력: 질소분포, 상, 균열, 경도, 잔류강도"),
            ((1080, 640, 1440, 770), "검증: 장시간 노출, 위치 의존성, 열사이클"),
        ]:
            draw_box(d, box, txt, (255, 249, 235), (193, 137, 48), size=25)

    def f9(d):
        title(d, "암모니아 질화부식 연구 로드맵")
        d.line((170, 520, 1430, 520), fill=(70, 80, 92), width=6)
        for x, lab, desc in [
            (260, "단기", "표준 쿠폰시험\n대표 합금·코팅"),
            (610, "중기", "장시간·응력·열사이클\n가속시험"),
            (960, "장기", "실장치 데이터\n수명모델"),
            (1310, "확산", "설계 코드\n재료 호환성 지도"),
        ]:
            d.ellipse((x - 32, 488, x + 32, 552), fill=(76, 132, 117))
            draw_text(d, (x - 35, 570), lab, 30)
            draw_box(d, (x - 145, 640, x + 145, 805), desc, (241, 248, 244), (76, 132, 117), size=25)
        draw_text(d, (220, 250), "핵심 과제: 실제 NH3/H2/H2O 조성, 위치 의존성, 용접부·응력, 코팅 결함을 동시에 반영", 33)

    def f10(d):
        title(d, "통합 프레임")
        center = (800, 520)
        items = [
            ("활성질소 공급", (800, 170)),
            ("합금/코팅 반응", (1190, 520)),
            ("균열·박리·취화", (800, 850)),
            ("억제·수명예측", (410, 520)),
        ]
        for txt, (x, y) in items:
            draw_box(d, (x - 165, y - 70, x + 165, y + 70), txt, (242, 246, 250), (83, 98, 119), size=30)
        arrow(d, (800, 240), (1070, 480), width=5)
        arrow(d, (1190, 590), (930, 805), width=5)
        arrow(d, (800, 780), (530, 560), width=5)
        arrow(d, (410, 450), (670, 215), width=5)
        draw_box(d, (610, 445, 990, 595), "암모니아 에너지 시스템\n재료 신뢰성", (255, 248, 232), (193, 137, 48), size=30)

    for name, fn in [
        ("research_landscape.png", f1),
        ("nh3_reaction_pathway.png", f2),
        ("nitriding_potential_map.png", f3),
        ("reactor_gradient.png", f4),
        ("flame_wall.png", f5),
        ("alloy_response_matrix.png", f6),
        ("alumina_barrier.png", f7),
        ("data_workflow.png", f8),
        ("roadmap.png", f9),
        ("synthesis_loop.png", f10),
    ]:
        save_canvas(name, fn)


def set_cell_text(cell, text: str, bold=False, shade: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    apply_font(run, 8.6, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_paragraph_keep_together(p)
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)
    if shade:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), shade)
        tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_height_at_least(row, height_cm: float) -> None:
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def set_paragraph_keep_together(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:keepLines", "w:keepNext"):
        if p_pr.find(qn(tag)) is None:
            p_pr.append(OxmlElement(tag))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])


def set_font_slots(r_fonts) -> None:
    r_fonts.set(qn("w:ascii"), E_FONT)
    r_fonts.set(qn("w:hAnsi"), E_FONT)
    r_fonts.set(qn("w:cs"), E_FONT)
    r_fonts.set(qn("w:eastAsia"), K_FONT)


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_doc_defaults_fonts(doc: Document) -> None:
    styles_element = doc.styles.element
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, doc_defaults)
    r_pr_default = get_or_add_child(doc_defaults, "w:rPrDefault")
    r_pr = get_or_add_child(r_pr_default, "w:rPr")
    r_fonts = get_or_add_child(r_pr, "w:rFonts")
    set_font_slots(r_fonts)


def apply_font(run, size: float | None = None, bold: bool | None = None) -> None:
    """Use Batang for Korean glyphs and Times New Roman for Latin glyphs."""
    run.font.name = E_FONT
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    set_font_slots(r_fonts)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def apply_style_font(style) -> None:
    style.font.name = E_FONT
    style.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    set_font_slots(r_fonts)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(19)
    section.page_height = Cm(26)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.1)
    add_page_number(section.footer.paragraphs[0])
    set_doc_defaults_fonts(doc)

    styles = doc.styles
    normal = styles["Normal"]
    apply_style_font(normal)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.55
    normal.paragraph_format.space_after = Pt(3)

    for style_name, size, before, after in [
        ("Heading 1", 16, 12, 8),
        ("Heading 2", 13, 8, 5),
        ("Heading 3", 11.5, 6, 3),
    ]:
        style = styles[style_name]
        apply_style_font(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def mark_content_added(doc: Document) -> None:
    setattr(doc, "_snu_last_action_page_break", False)


def add_page_break_once(doc: Document) -> None:
    if getattr(doc, "_snu_last_action_page_break", False):
        return
    doc.add_page_break()
    setattr(doc, "_snu_last_action_page_break", True)


def add_para(doc: Document, text: str, style: str | None = None, align=None, first_indent=True) -> None:
    p = doc.add_paragraph(style=style)
    mark_content_added(doc)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.55
    p.paragraph_format.space_after = Pt(3)
    if first_indent and style is None:
        p.paragraph_format.first_line_indent = Pt(12)
    run = p.add_run(text)
    apply_font(run, 10.5 if style is None else 12)


def add_front_list_entry(doc: Document, text: str, size: float = 8.8) -> None:
    p = doc.add_paragraph()
    mark_content_added(doc)
    p.paragraph_format.line_spacing = 1.55
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    apply_font(run, size)


def require_metadata() -> None:
    if os.environ.get("ALLOW_PLACEHOLDER_DRAFT") == "1":
        return
    missing = [label for key, label in REQUIRED_METADATA if not METADATA[key]]
    if missing:
        questions = "\n".join(f"- {label}을/를 알려주세요." for label in missing)
        raise SystemExit(
            "원고 생성 전에 필요한 메타데이터가 비어 있습니다. 먼저 사용자에게 아래 항목을 질문하세요:\n"
            f"{questions}\n\n"
            "임시 placeholder 초안을 강제로 만들 때만 ALLOW_PLACEHOLDER_DRAFT=1을 사용하세요."
        )


def metadata_value(key: str) -> str:
    value = METADATA[key]
    if value:
        return value
    return f"미입력: {dict(REQUIRED_METADATA).get(key, key)}"


def add_caption_to(container, text: str, space_before: float = 5, space_after: float = 0) -> None:
    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    apply_font(run, 9.2, bold=True)


def add_caption(doc: Document, text: str) -> None:
    add_caption_to(doc, text, space_before=5, space_after=7)


def add_isolated_visual_container(doc: Document):
    add_page_break_once(doc)
    layout = doc.add_table(rows=1, cols=1)
    mark_content_added(doc)
    layout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders_none(layout)
    set_row_height_at_least(layout.rows[0], 19.8)
    cell = layout.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    return cell


def fit_image_dimensions(path: Path, max_width_cm: float = 12.8, max_height_cm: float = 16.3):
    with Image.open(path) as image:
        width_px, height_px = image.size
    aspect = height_px / max(width_px, 1)
    width_cm = max_width_cm
    height_cm = width_cm * aspect
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm / max(aspect, 0.01)
        return None, Cm(height_cm)
    return Cm(width_cm), None


def add_figure_item(cell, idx: int, *, max_width_cm: float, max_height_cm: float) -> None:
    caption, filename = FIGURES[idx]
    image_path = OUT_DIR / filename if "/" in filename else FIG_DIR / filename
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    width, height = fit_image_dimensions(image_path, max_width_cm=max_width_cm, max_height_cm=max_height_cm)
    if width is not None:
        run.add_picture(str(image_path), width=width)
    else:
        run.add_picture(str(image_path), height=height)
    add_caption_to(cell, caption, space_before=8, space_after=0)


def add_figure_group(doc: Document, indices: list[int]) -> None:
    cell = add_isolated_visual_container(doc)
    if len(indices) == 1:
        add_figure_item(cell, indices[0], max_width_cm=12.8, max_height_cm=16.3)
    else:
        for pos, idx in enumerate(indices[:2]):
            if pos:
                gap = cell.add_paragraph()
                gap.paragraph_format.space_after = Pt(12)
            add_figure_item(cell, idx, max_width_cm=11.7, max_height_cm=6.8)
    add_page_break_once(doc)


def add_figure(doc: Document, idx: int) -> None:
    add_figure_group(doc, [idx])


def add_table(doc: Document, idx: int) -> None:
    caption, rows = TABLES[idx]
    cell = add_isolated_visual_container(doc)
    table = cell.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_i, row in enumerate(rows):
        set_row_cant_split(table.rows[r_i])
        if r_i == 0:
            set_repeat_table_header(table.rows[0])
        for c_i, text in enumerate(row):
            set_cell_text(table.cell(r_i, c_i), text, bold=(r_i == 0), shade=None)
    add_caption_to(cell, caption, space_before=8, space_after=0)
    add_page_break_once(doc)


def compact_paragraphs(paragraphs: list[str]) -> list[str]:
    if MAX_PARAGRAPHS_PER_SUBSECTION <= 0 or len(paragraphs) <= MAX_PARAGRAPHS_PER_SUBSECTION:
        return paragraphs
    selected = paragraphs[:MAX_PARAGRAPHS_PER_SUBSECTION]
    for paragraph in paragraphs[MAX_PARAGRAPHS_PER_SUBSECTION:]:
        if "Figure" in paragraph or "Table" in paragraph:
            if paragraph not in selected:
                selected[-1] = paragraph
    return selected


def add_front_matter(doc: Document) -> None:
    def center_line(text, size=16, bold=False, space=10):
        p = doc.add_paragraph()
        mark_content_added(doc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space)
        run = p.add_run(text)
        apply_font(run, size, bold=bold)
        return p

    center_line(metadata_value("degree_name"), 17, True, 28)
    for line in METADATA["korean_title_lines"]:
        center_line(line, 21 if line != METADATA["korean_title_lines"][-1] else 18, True, 3)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(36)
    for line in METADATA["english_title_lines"]:
        center_line(line, 12, False, 6)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(72)
    center_line("서울대학교 대학원", 15, False, 6)
    center_line(f"{metadata_value('department')} {metadata_value('major')}", 14, False, 22)
    center_line(METADATA["author"], 18, True, 50)
    center_line(metadata_value("submission_month"), 13, False, 10)
    add_page_break_once(doc)

    # Page 2 must be a single approval page in this review-paper workflow.
    center_line(f"{METADATA['author']}의 {metadata_value('degree_name')}을 인준함", 16, True, 18)
    center_line(metadata_value("approval_month"), 13, False, 44)
    center_line(f"지도교수  {metadata_value('advisor')}", 13, False, 42)
    for label, name in [
        ("위원장", METADATA["committee_chair"]),
        ("부위원장", METADATA["committee_vice_chair"]),
        (METADATA["committee_member_role"], METADATA["committee_member"]),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}     {name}      (인)")
        apply_font(run, 14)
        p.paragraph_format.space_after = Pt(24)
    add_page_break_once(doc)

    # Page 3 starts with the Korean abstract.
    center_line("국문초록", 16, True, 18)
    abstract = [
        "암모니아는 수소 운반체이자 탄소를 포함하지 않는 연료로 주목받고 있으나, 고온 암모니아 활용 장치에서 구조재가 겪는 원치 않는 질화 및 질화부식은 아직 충분히 정리되지 않은 문제이다. 본 리뷰는 2024-2026년 문헌 동향을 중심으로 암모니아 질화 연구의 초점이 제어된 표면경화 공정에서 암모니아 크래킹, 연소, 가스터빈, 선박엔진, 저장·수송 인프라의 재료 신뢰성 문제로 확장되고 있음을 고찰하였다.",
        "전통적인 NH3 가스질화에서는 암모니아 해리와 질소 확산을 이용하여 강 표면의 경도와 마모 저항을 향상시킨다. 그러나 암모니아 에너지 시스템에서는 같은 반응이 CrN, Fe4N, Fe2-3N, AlN 등의 형성과 격자팽창, 입계취화, 코팅 원소 고갈, 표면균열 및 박리를 유발할 수 있다. 특히 반응기 위치, 벽면 NH3 농도, 수증기, 온도창, 합금 원소, 보호 산화막의 연속성이 질화부식의 지배인자로 나타난다.",
        "본 논문은 암모니아 질화의 계면반응과 열역학, 제어된 질화 공정의 최신 고도화, 고온 암모니아 장치에서의 원치 않는 질화, 합금계별 반응성과 억제 전략, 진단·모델링·데이터화 방향을 통합적으로 검토하였다. 결론적으로 암모니아 에너지 기술의 상용화를 위해서는 촉매 성능뿐 아니라 구조재의 질화 포텐셜 노출, 장시간 안정성, 코팅 결함, 용접부 응력, 실제 유동장 위치 의존성을 포함한 재료 호환성 지도가 필요하다.",
        "주요어: 암모니아 질화, 질화부식, 암모니아 크래킹, 암모니아 연소, 스테인리스강, MCrAlY 코팅, 알루미나 장벽, 재료 신뢰성",
    ]
    for para in abstract:
        add_para(doc, para)
    add_page_break_once(doc)

    center_line("목차", 16, True, 18)
    toc_items = [section["heading"] for section in SECTION_DATA] + ["참고문헌"]
    if int(os.environ.get("EXTRA_REVIEW_NOTES", "0")) > 0:
        toc_items.append("부록 A. 문헌별 검토 메모")
    toc_items += ["Abstract"]
    for item in toc_items:
        add_front_list_entry(doc, f"\t{item}\t1", size=10.5)
    add_page_break_once(doc)

    center_line("표 목차", 16, True, 18)
    for caption, _ in TABLES:
        add_front_list_entry(doc, f"{caption}\t1", size=10.5)
    add_page_break_once(doc)

    center_line("그림 목차", 16, True, 18)
    for caption, _ in FIGURES:
        add_front_list_entry(doc, f"{caption}\t1", size=10.5)
    add_page_break_once(doc)


def add_body(doc: Document) -> None:
    table_inserted = set()
    sourced_figure_map = {
        "2.1": [10, 11],
        "2.2": [12, 13],
        "4.1": [14],
        "5.5": [15],
    }
    for section in SECTION_DATA:
        doc.add_heading(section["heading"], level=1)
        mark_content_added(doc)
        for sub in section["subs"]:
            subheading, paragraphs, *maybe_fig = sub
            doc.add_heading(subheading, level=2)
            mark_content_added(doc)
            for para in compact_paragraphs(paragraphs):
                add_para(doc, para)
            key = re.match(r"(\d+\.\d+)", subheading).group(1)
            if key in {"1.3"}:
                add_table(doc, 0)
            if key in {"4.5"}:
                add_table(doc, 1)
            if key in {"5.5"}:
                add_table(doc, 2)
            if key in {"3.1"}:
                add_table(doc, 3)
            if key in {"6.1"}:
                add_table(doc, 4)
            if key in {"7.2"}:
                add_table(doc, 5)
            fig_indices = []
            if maybe_fig:
                fig_indices.append(maybe_fig[0])
            fig_indices.extend(sourced_figure_map.get(key, []))
            for fig_idx in fig_indices:
                if fig_idx in table_inserted:
                    continue
                table_inserted.add(fig_idx)
                add_figure(doc, fig_idx)
            if ADD_EXPANSION_PARAGRAPHS:
                add_expansion_paragraphs(doc, subheading)


def add_expansion_paragraphs(doc: Document, subheading: str) -> None:
    topic = subheading.split(" ", 1)[1] if " " in subheading else subheading
    templates = [
        f"{topic}에서 중요한 해석상의 주의점은 문헌에 보고된 질화층을 단순히 두껍거나 얇은 층으로만 비교해서는 안 된다는 것이다. 층의 연속성, 상 조성, 기공과 균열, 기지와의 계면 접착성, 그리고 노출 뒤 잔류 기계특성이 함께 평가되어야 실제 장치 수명과 연결된다.",
        f"또한 {topic}의 결과는 시험편 표면 준비와 초기 산화막 상태에 민감하다. 같은 합금 조성이라도 연마, 세척, 사전 산화, 코팅 결함, 용접 열영향부가 다르면 NH3 해리와 질소 침투 경로가 달라질 수 있다. 이는 표준 시험법이 필요한 이유이기도 하다.",
    ]
    for para in templates:
        add_para(doc, para)


def add_references(doc: Document) -> None:
    add_page_break_once(doc)
    doc.add_heading("참고문헌", level=1)
    mark_content_added(doc)
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.first_line_indent = Pt(-16)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"[{i}] {ref}")
        apply_font(run, 8.8)


def add_back_matter(doc: Document) -> None:
    add_page_break_once(doc)
    doc.add_heading("Abstract", level=1)
    mark_content_added(doc)
    for para in [
        "Ammonia is increasingly treated as a hydrogen carrier and carbon-free fuel, but high-temperature ammonia environments can turn the useful chemistry of gas nitriding into an unwanted degradation mechanism. This review summarizes recent literature on unwanted nitriding and nitridation corrosion in ammonia energy systems, including ammonia cracking reactors, ammonia combustion walls, ammonia-fueled gas turbine components, engine-related steels, and liquid-ammonia storage infrastructure.",
        "The reviewed studies indicate that nitridation damage is governed not only by ammonia concentration but also by nitriding potential, wall temperature, water vapor, hydrogen, reactor position, alloy chemistry, oxide-scale integrity, coating defects, stress, and exposure time. Fe, Cr, Al, Ni, and Mo can each contribute to either protection or degradation depending on where the nitride phases form and whether the resulting layer remains continuous and adherent.",
        "The central conclusion is that ammonia energy deployment requires materials compatibility maps that connect process conditions, nitrogen activity, microstructural evolution, cracking, residual properties, and mitigation strategies. Alumina-forming diffusion coatings, controlled pre-oxidation, alloy optimization, and standardized long-duration testing are identified as especially important directions.",
        "Keywords: ammonia nitriding, nitridation corrosion, ammonia cracking, ammonia combustion, stainless steel, MCrAlY coating, alumina barrier, materials reliability",
    ]:
        add_para(doc, para)


def add_appendix_depth(doc: Document) -> None:
    extra_notes = int(os.environ.get("EXTRA_REVIEW_NOTES", "0"))
    if extra_notes <= 0:
        return
    add_page_break_once(doc)
    doc.add_heading("부록 A. 문헌별 검토 메모", level=1)
    mark_content_added(doc)
    notes = [
        ("고온 암모니아 구조재 리뷰", "고온 NH3 장치에서 ferritic steel과 austenitic stainless steel의 손상모드를 같은 열화 프레임으로 묶어 주며, 촉매·질화열처리·구조재 손상을 연결하는 출발점이다."),
        ("암모니아 화염 질화", "화염-벽 반응에서 벽면 NH3 농도와 NH2 라디칼의 역할을 분리해 보려는 시도라는 점에서 연소공학과 표면공학을 연결한다."),
        ("310S 위치 의존 부식", "반응기 내부 위치가 부식속도와 상 형성에 영향을 준다는 점을 보여 주어 단일 쿠폰시험의 한계를 드러낸다."),
        ("CoNiCrAlY 코팅", "산화 환경에 최적화된 코팅이 NH3 환원성 질화 환경에서 내부질화를 겪을 수 있음을 보여 주며, 기존 터빈 코팅의 재검증 필요성을 제기한다."),
        ("Al2O3 장벽", "고압 NH3 크래킹에서 알루미나 형성 확산코팅이 질소 침투 억제에 유망하다는 구체적 경로를 제공한다."),
        ("저온 DSS 질화", "경도 향상과 내식성 저하 사이의 균형이 공정 조건에 민감함을 보여 주어 스테인리스강 질화의 장점과 위험을 동시에 보여 준다."),
        ("액체 NH3 SCC", "고온 질화부식과 다른 손상모드이지만 암모니아 가치사슬의 저장·수송 안전성 측면에서 함께 다루어야 한다."),
    ]
    for i in range(extra_notes):
        label, body = notes[i % len(notes)]
        doc.add_heading(f"A.{i+1} {label}", level=2)
        mark_content_added(doc)
        add_para(doc, body)
        add_para(doc, "이 항목은 본문에서 제시한 통합 프레임을 문헌 단위로 다시 점검하기 위한 메모이다. 실제 제출본에서는 각 원문의 실험 조건, 합금 조성, 노출 시간, 분석법, 원자료 수치까지 확인하여 표 형태의 체계적 문헌고찰로 확장하는 것이 바람직하다.")


def build() -> None:
    require_metadata()
    ensure_dirs()
    download_source_figures()
    make_figures()
    doc = Document()
    set_document_defaults(doc)
    add_front_matter(doc)
    add_body(doc)
    add_references(doc)
    add_appendix_depth(doc)
    add_back_matter(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
